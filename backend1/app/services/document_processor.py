# app/services/document_processor.py
# Service xử lý tài liệu PDF và tạo chunks cho RAG

import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
import pytesseract
import cv2
import numpy as np
from PIL import Image
import docx  # python-docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from app.core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Class xử lý tài liệu PDF và tạo chunks"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Cấu hình OCR
        self.use_ocr = settings.USE_OCR
        if self.use_ocr:
            self._setup_ocr()
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Trích xuất text từ file PDF với hỗ trợ OCR"""
        try:
            text = ""
            
            # Thử với PyMuPDF trước (nhanh hơn)
            try:
                doc = fitz.open(pdf_path)
                for page in doc:
                    page_text = page.get_text()
                    if page_text.strip():
                        text += page_text
                doc.close()
                if text.strip():
                    logger.info(f"Đã trích xuất text từ {pdf_path} bằng PyMuPDF")
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {pdf_path}: {e}")
            
            # Thử với pdfplumber
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    logger.info(f"Đã trích xuất text từ {pdf_path} bằng pdfplumber")
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber failed for {pdf_path}: {e}")
            
            # Thử với PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    logger.info(f"Đã trích xuất text từ {pdf_path} bằng PyPDF2")
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {pdf_path}: {e}")
            
            # Nếu tất cả phương pháp thông thường thất bại, thử OCR
            if self.use_ocr:
                logger.info(f"Thử OCR cho {pdf_path}")
                ocr_text = self._extract_text_with_ocr(pdf_path)
                if ocr_text.strip():
                    logger.info(f"Đã trích xuất text từ {pdf_path} bằng OCR")
                    return ocr_text
            
            logger.error(f"Không thể trích xuất text từ {pdf_path}")
            return ""
            
        except Exception as e:
            logger.error(f"Lỗi khi xử lý file {pdf_path}: {e}")
            return ""
    
    def _setup_ocr(self):
        """Thiết lập cấu hình OCR"""
        try:
            # Cấu hình đường dẫn tesseract
            if os.path.exists(settings.TESSERACT_PATH):
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
                logger.info(f"Đã cấu hình Tesseract tại: {settings.TESSERACT_PATH}")
            else:
                logger.warning(f"Không tìm thấy Tesseract tại: {settings.TESSERACT_PATH}")
                logger.info("Sử dụng Tesseract từ PATH")
            
            # Test OCR
            test_image = Image.new('RGB', (100, 100), color='white')
            pytesseract.image_to_string(test_image, lang=settings.OCR_LANGUAGES)
            logger.info("OCR đã được thiết lập thành công")
            
        except Exception as e:
            logger.error(f"Lỗi khi thiết lập OCR: {e}")
            self.use_ocr = False
    
    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Tiền xử lý hình ảnh để cải thiện OCR"""
        try:
            # Chuyển đổi sang grayscale
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
            
            # Làm mờ để giảm noise
            blurred = cv2.GaussianBlur(gray, (3, 3), 0)
            
            # Tăng cường độ tương phản
            enhanced = cv2.convertScaleAbs(blurred, alpha=1.2, beta=10)
            
            # Threshold để tạo ảnh nhị phân
            _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            return binary
            
        except Exception as e:
            logger.warning(f"Lỗi khi tiền xử lý hình ảnh: {e}")
            return image
    
    def _extract_text_with_ocr(self, pdf_path: str) -> str:
        """Trích xuất text từ PDF sử dụng OCR"""
        try:
            if not self.use_ocr:
                return ""
            
            text = ""
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Chuyển đổi trang thành hình ảnh
                mat = fitz.Matrix(2.0, 2.0)  # Tăng độ phân giải
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Chuyển đổi thành numpy array
                nparr = np.frombuffer(img_data, np.uint8)
                image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                
                # Tiền xử lý hình ảnh
                processed_image = self._preprocess_image(image)
                
                # OCR
                page_text = pytesseract.image_to_string(
                    processed_image,
                    lang=settings.OCR_LANGUAGES,
                    config=f'--psm {settings.OCR_PSM}'
                )
                
                if page_text.strip():
                    text += page_text + "\n"
                    logger.info(f"Đã OCR trang {page_num + 1} của {pdf_path}")
            
            doc.close()
            return text
            
        except Exception as e:
            logger.error(f"Lỗi khi OCR file {pdf_path}: {e}")
            return ""
    
    def extract_text_from_word(self, word_path: str) -> str:
        """Trích xuất text từ file Word (.docx)"""
        try:
            doc = docx.Document(word_path)
            text = ""
            
            # Trích xuất text từ paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Trích xuất text từ tables (nếu có)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            if text.strip():
                logger.info(f"Đã trích xuất text từ Word document: {word_path}")
                return text
            else:
                logger.warning(f"Không tìm thấy text trong Word document: {word_path}")
                return ""
                
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất text từ Word document {word_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """Trích xuất text từ file .txt với nhiều encoding"""
        try:
            encodings = ['utf-8', 'utf-16', 'cp1252', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Đã đọc file TXT với encoding {encoding}: {txt_path}")
                    return text
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Không thể đọc file {txt_path} với các encoding đã thử")
            return ""
            
        except Exception as e:
            logger.error(f"Lỗi đọc TXT cho file {txt_path}: {e}")
            return ""
    
    def extract_text_from_document(self, doc_path: str) -> str:
        """Trích xuất text từ các loại file tài liệu khác nhau"""
        try:
            file_ext = Path(doc_path).suffix.lower()
            
            if file_ext == '.pdf':
                return self.extract_text_from_pdf(doc_path)
            elif file_ext in ['.doc', '.docx']:
                return self.extract_text_from_word(doc_path)
            elif file_ext == '.txt':
                return self.extract_text_from_txt(doc_path)
            else:
                logger.warning(f"Loại file không được hỗ trợ: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"Lỗi xử lý file {doc_path}: {e}")
            return ""
    
    def extract_text_from_image(self, image_path: str) -> str:
        """Trích xuất text từ hình ảnh sử dụng OCR"""
        try:
            if not self.use_ocr:
                logger.warning("OCR không được kích hoạt")
                return ""
            
            # Đọc hình ảnh
            image = cv2.imread(image_path)
            if image is None:
                logger.error(f"Không thể đọc hình ảnh: {image_path}")
                return ""
            
            # Tiền xử lý hình ảnh
            processed_image = self._preprocess_image(image)
            
            # OCR
            text = pytesseract.image_to_string(
                processed_image,
                lang=settings.OCR_LANGUAGES,
                config=f'--psm {settings.OCR_PSM}'
            )
            
            if text.strip():
                logger.info(f"Đã OCR hình ảnh: {image_path}")
                return text
            else:
                logger.warning(f"Không tìm thấy text trong hình ảnh: {image_path}")
                return ""
                
        except Exception as e:
            logger.error(f"Lỗi khi OCR hình ảnh {image_path}: {e}")
            return ""
    
    def process_image_directory(self, directory_path: str, category: str = "images") -> List[Document]:
        """Xử lý tất cả hình ảnh trong thư mục"""
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Thư mục không tồn tại: {directory_path}")
            return documents
        
        # Các định dạng hình ảnh được hỗ trợ
        image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
        image_files = []
        
        for ext in image_extensions:
            image_files.extend(directory.glob(f"*{ext}"))
            image_files.extend(directory.glob(f"*{ext.upper()}"))
        
        logger.info(f"Tìm thấy {len(image_files)} file hình ảnh trong {directory_path}")
        
        for image_file in image_files:
            try:
                # Trích xuất text từ hình ảnh
                text = self.extract_text_from_image(str(image_file))
                if not text:
                    continue
                
                # Làm sạch text
                text = self.clean_text(text)
                
                # Tạo metadata
                metadata = {
                    "source": str(image_file),
                    "filename": image_file.name,
                    "category": category,
                    "file_type": "image"
                }
                
                # Tạo document
                doc = Document(page_content=text, metadata=metadata)
                
                # Chia nhỏ thành chunks
                chunks = self.text_splitter.split_documents([doc])
                
                # Thêm thông tin chunk vào metadata
                for i, chunk in enumerate(chunks):
                    chunk.metadata.update({
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    })
                
                documents.extend(chunks)
                logger.info(f"Đã xử lý hình ảnh: {image_file.name}")
                
            except Exception as e:
                logger.error(f"Lỗi khi xử lý hình ảnh {image_file}: {e}")
                continue
        
        logger.info(f"Tổng cộng xử lý được {len(documents)} documents từ hình ảnh")
        return documents
    
    def clean_text(self, text: str) -> str:
        """Làm sạch text"""
        # Loại bỏ ký tự đặc biệt và chuẩn hóa
        import re
        
        # Loại bỏ ký tự điều khiển
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Chuẩn hóa khoảng trắng
        text = re.sub(r'\s+', ' ', text)
        
        # Loại bỏ dòng trống
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text.strip()
    
    def create_documents_from_pdf(self, pdf_path: str, category: str = "general") -> List[Document]:
        """Tạo documents từ file PDF"""
        try:
            # Trích xuất text
            text = self.extract_text_from_pdf(pdf_path)
            if not text:
                return []
            
            # Làm sạch text
            text = self.clean_text(text)
            
            # Tạo metadata
            filename = Path(pdf_path).name
            metadata = {
                "source": pdf_path,
                "filename": filename,
                "category": category,
                "file_type": "pdf"
            }
            
            # Tạo document
            doc = Document(page_content=text, metadata=metadata)
            
            # Chia nhỏ thành chunks
            chunks = self.text_splitter.split_documents([doc])
            
            # Thêm thông tin chunk vào metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                })
            
            logger.info(f"Đã tạo {len(chunks)} chunks từ {filename}")
            return chunks
            
        except Exception as e:
            logger.error(f"Lỗi khi tạo documents từ {pdf_path}: {e}")
            return []
    
    def create_documents_from_file(self, file_path: str, category: str = "general") -> List[Document]:
        """Tạo documents từ file bất kỳ (PDF, Word, TXT)"""
        try:
            # Trích xuất text
            text = self.extract_text_from_document(file_path)
            if not text:
                return []
            
            # Làm sạch text
            text = self.clean_text(text)
            
            # Tạo metadata
            filename = Path(file_path).name
            file_ext = Path(file_path).suffix.lower()
            
            metadata = {
                "source": file_path,
                "filename": filename,
                "category": category,
                "file_type": file_ext.lstrip('.')
            }
            
            # Tạo document
            doc = Document(page_content=text, metadata=metadata)
            
            # Chia nhỏ thành chunks
            chunks = self.text_splitter.split_documents([doc])
            
            # Thêm thông tin chunk vào metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                })
            
            logger.info(f"Đã tạo {len(chunks)} chunks từ {filename}")
            return chunks
            
        except Exception as e:
            logger.error(f"Lỗi khi tạo documents từ {file_path}: {e}")
            return []
    
    def process_directory(self, directory_path: str, category: str = "general") -> List[Document]:
        """Xử lý tất cả file tài liệu trong thư mục (PDF, Word, TXT)"""
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Thư mục không tồn tại: {directory_path}")
            return documents
        
        # Các định dạng file được hỗ trợ
        supported_extensions = ['.pdf', '.doc', '.docx', '.txt']
        document_files = []
        
        for ext in supported_extensions:
            document_files.extend(directory.glob(f"*{ext}"))
            document_files.extend(directory.glob(f"*{ext.upper()}"))
        
        logger.info(f"Tìm thấy {len(document_files)} file tài liệu trong {directory_path}")
        
        for doc_file in document_files:
            try:
                chunks = self.create_documents_from_file(str(doc_file), category)
                documents.extend(chunks)
            except Exception as e:
                logger.error(f"Lỗi khi xử lý {doc_file}: {e}")
                continue
        
        logger.info(f"Tổng cộng tạo được {len(documents)} documents từ {directory_path}")
        return documents
    
    def process_all_document_directories(self, luat_path: str, english_path: str, vietnamese_path: str) -> List[Document]:
        """Xử lý tất cả thư mục tài liệu (PDF và hình ảnh)"""
        all_documents = []
        
        # Xử lý tài liệu luật
        if os.path.exists(luat_path):
            luat_docs = self.process_directory(luat_path, "luat")
            all_documents.extend(luat_docs)
            
            # Xử lý hình ảnh trong thư mục luật (nếu có)
            luat_images = self.process_image_directory(luat_path, "luat_images")
            all_documents.extend(luat_images)
        
        # Xử lý tài liệu tiếng Anh
        if os.path.exists(english_path):
            english_docs = self.process_directory(english_path, "english")
            all_documents.extend(english_docs)
            
            # Xử lý hình ảnh trong thư mục tiếng Anh (nếu có)
            english_images = self.process_image_directory(english_path, "english_images")
            all_documents.extend(english_images)
        
        # Xử lý tài liệu tiếng Việt
        if os.path.exists(vietnamese_path):
            vietnamese_docs = self.process_directory(vietnamese_path, "vietnamese")
            all_documents.extend(vietnamese_docs)
            
            # Xử lý hình ảnh trong thư mục tiếng Việt (nếu có)
            vietnamese_images = self.process_image_directory(vietnamese_path, "vietnamese_images")
            all_documents.extend(vietnamese_images)
        
        logger.info(f"Tổng cộng xử lý được {len(all_documents)} documents (PDF + hình ảnh)")
        return all_documents
